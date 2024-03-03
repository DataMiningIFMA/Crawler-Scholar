import csv
from docx import Document

def converter_docx_em_csv(caminho_arquivo_docx, caminho_arquivo_csv):
    doc = Document('c:\\Users\\there\\Documents\\Crawler-Scholar\\linksperf.docx')
    linhas = []
# Iterar pelos parágrafos e tabelas no arquivo docx
    for paragrafo in doc.paragraphs:
        linhas.append(paragrafo.text.split('\n'))

    for tabela in doc.tables:
        for linha in tabela.rows:
            linhas.append([celula.text for celula in linha.cells])

    with open('c:\\Users\\there\\Documents\\Crawler-Scholar\\perfs_csv.csv', 'w', newline='') as arquivo_csv:
        escritor = csv.writer(arquivo_csv)
        escritor.writerows(linhas)

converter_docx_em_csv('c:\\Users\\there\\Documents\\Crawler-Scholar\\linksperf.docx', 'c:\\Users\\there\\Documents\\Crawler-Scholar\\perfs.csv')